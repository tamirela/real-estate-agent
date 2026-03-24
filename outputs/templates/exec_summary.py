"""
Executive Summary document generator.
Creates a Word docx with property overview, financial snapshot,
valuation analysis, risk assessment, value-add opportunities,
and GO/NO-GO verdict.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _safe(d, key, default=None):
    v = d.get(key, default)
    return v if v is not None else default


def _fmt_dollar(val):
    try:
        return f"${float(val):,.0f}"
    except (ValueError, TypeError):
        return str(val) if val else "N/A"


def _fmt_pct(val):
    try:
        return f"{float(val):.1%}"
    except (ValueError, TypeError):
        return str(val) if val else "N/A"


def _set_cell_font(cell, text, bold=False, size=10):
    """Set cell text with Calibri font."""
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold


def _add_heading(doc, text, level=1):
    """Add a heading with Calibri font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return heading


def _add_table_with_borders(doc, rows, cols):
    """Create a table with thin borders."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply borders via XML
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(
            qn(f"w:{edge}"),
            {"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": "999999"},
        )
        borders.append(element)
    tbl_pr.append(borders)

    return table


# ─────────────────────────────────────────────────────────────────
# SECTIONS
# ─────────────────────────────────────────────────────────────────

def _section_property_overview(doc, d):
    _add_heading(doc, "1. Property Overview", level=1)

    info = [
        ("Property Name", _safe(d, "property_name", "N/A")),
        ("Address", _safe(d, "address", "N/A")),
        ("City, State", f"{_safe(d, 'city', '')}, {_safe(d, 'state', 'TX')}"),
        ("Units", _safe(d, "units", "N/A")),
        ("Total SF", f"{_safe(d, 'sqft', 'N/A'):,}" if isinstance(_safe(d, "sqft"), (int, float)) else "N/A"),
        ("Year Built", _safe(d, "year_built", "N/A")),
        ("Asking Price", _fmt_dollar(_safe(d, "price"))),
        ("Price Per Unit", _fmt_dollar(_safe(d, "price_per_unit"))),
        ("Source", _safe(d, "source", "N/A")),
    ]

    table = _add_table_with_borders(doc, len(info), 2)
    for i, (label, value) in enumerate(info):
        _set_cell_font(table.rows[i].cells[0], label, bold=True)
        _set_cell_font(table.rows[i].cells[1], value)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.0)


def _section_financial_snapshot(doc, d):
    _add_heading(doc, "2. Financial Snapshot", level=1)

    noi = _safe(d, "noi", 0)
    va_noi = _safe(d, "va_noi", 0)

    metrics = [
        ("Metric", "T-12 (As-Is)", "Pro Forma (Stabilized)"),
        ("Net Operating Income", _fmt_dollar(noi), _fmt_dollar(va_noi)),
        ("Cap Rate", _fmt_pct(_safe(d, "cap_rate")), _fmt_pct(_safe(d, "va_cap_rate"))),
        ("Cash-on-Cash Return", _fmt_pct(_safe(d, "cash_on_cash")), _fmt_pct(_safe(d, "va_cash_on_cash"))),
        ("DSCR", f"{_safe(d, 'dscr', 0):.2f}", ""),
        ("GRM", f"{_safe(d, 'grm', 0):.1f}", ""),
        ("Annual Cash Flow", _fmt_dollar(_safe(d, "annual_cash_flow")), ""),
        ("5-Year IRR", _fmt_pct(_safe(d, "irr_5yr")), ""),
        ("Equity Multiple (5yr)", f"{_safe(d, 'equity_multiple_5yr', 0):.2f}x", ""),
        ("Exit Value (5yr)", _fmt_dollar(_safe(d, "exit_value")), ""),
    ]

    table = _add_table_with_borders(doc, len(metrics), 3)
    for i, (col1, col2, col3) in enumerate(metrics):
        is_header = (i == 0)
        _set_cell_font(table.rows[i].cells[0], col1, bold=True)
        _set_cell_font(table.rows[i].cells[1], col2, bold=is_header)
        _set_cell_font(table.rows[i].cells[2], col3, bold=is_header)


def _section_valuation(doc, d):
    _add_heading(doc, "3. Valuation at Different Cap Rates", level=1)

    noi = _safe(d, "noi", 0)
    va_noi = _safe(d, "va_noi", 0)

    cap_rates = [0.055, 0.060, 0.065, 0.070, 0.075, 0.080]

    headers = ["Cap Rate", "Value (T-12 NOI)", "Value (Stabilized NOI)"]
    table = _add_table_with_borders(doc, len(cap_rates) + 1, 3)

    for c, hdr in enumerate(headers):
        _set_cell_font(table.rows[0].cells[c], hdr, bold=True)

    for i, cap in enumerate(cap_rates):
        row = table.rows[i + 1]
        _set_cell_font(row.cells[0], f"{cap:.1%}", bold=True)
        val_t12 = noi / cap if cap > 0 else 0
        val_va = va_noi / cap if cap > 0 else 0
        _set_cell_font(row.cells[1], _fmt_dollar(val_t12))
        _set_cell_font(row.cells[2], _fmt_dollar(val_va))


def _section_key_risks(doc, d):
    _add_heading(doc, "4. Key Risks", level=1)

    red_flags = _safe(d, "red_flags", [])
    if red_flags:
        for flag in red_flags:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(flag)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
    else:
        p = doc.add_paragraph("No significant red flags identified.")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)


def _section_value_add(doc, d):
    _add_heading(doc, "5. Value-Add Opportunities", level=1)

    signals = _safe(d, "value_add_signals", [])
    if signals:
        for signal in signals:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(signal)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
    else:
        p = doc.add_paragraph("No specific value-add signals identified from listing data.")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)

    # Renovation assumptions
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Renovation Assumptions: ")
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(10)
    units = _safe(d, "units", 0)
    reno_per_unit = _safe(d, "reno_cost_per_unit", 8000)
    run2 = p.add_run(
        f"${reno_per_unit:,}/unit x {units} units = "
        f"${reno_per_unit * units:,} total renovation budget. "
        f"Target {_fmt_pct(_safe(d, 'value_add_rent_bump_pct', 0.20))} rent increase post-renovation."
    )
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)


def _section_verdict(doc, d):
    _add_heading(doc, "6. Preliminary GO / NO-GO", level=1)

    verdict = _safe(d, "verdict", "")
    reason = _safe(d, "hurdle_reason", _safe(d, "reason", ""))
    passes = _safe(d, "passes_hurdle", False)

    # Verdict line
    p = doc.add_paragraph()
    run = p.add_run(f"VERDICT:  {verdict if verdict else ('GO' if passes else 'NO-GO')}")
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if passes else RGBColor(0xE7, 0x4C, 0x3C)

    if reason:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(reason)
        run2.font.name = "Calibri"
        run2.font.size = Pt(10)

    # Disclaimer
    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    run3 = p3.add_run(
        "This analysis is preliminary and based on available listing data. "
        "Actual financials, property condition, and market conditions must be "
        "verified through due diligence before making any investment decision."
    )
    run3.font.name = "Calibri"
    run3.font.size = Pt(9)
    run3.font.italic = True
    run3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


# ─────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────

def generate(deal_data: dict, filepath: str) -> str:
    """
    Generate an Executive Summary docx and save it to *filepath*.

    Parameters
    ----------
    deal_data : dict
        Property and financial data for the deal.
    filepath : str
        Destination path for the .docx file.

    Returns
    -------
    str  The filepath that was written.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Title
    title = doc.add_heading("Executive Summary", level=0)
    for run in title.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    property_name = _safe(deal_data, "property_name", "")
    if property_name:
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(property_name)
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Build all sections
    _section_property_overview(doc, deal_data)
    _section_financial_snapshot(doc, deal_data)
    _section_valuation(doc, deal_data)
    _section_key_risks(doc, deal_data)
    _section_value_add(doc, deal_data)
    _section_verdict(doc, deal_data)

    doc.save(filepath)
    return filepath
